"""
Loader de DOCX — extrai texto, tabelas e metadados.
"""

from docx import Document
from . import BaseLoader, ExtractedContent


class DOCXLoader(BaseLoader):

    def extract(self, file_path: str) -> ExtractedContent:
        doc = Document(file_path)
        text_parts = []
        metadata = {
            "paragraphs": 0,
            "tables": 0,
        }

        # --- Propriedades do documento ---
        props = doc.core_properties
        if props.title:
            metadata["title"] = props.title
        if props.author:
            metadata["author"] = props.author

        # --- Parágrafos ---
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            metadata["paragraphs"] += 1

            # Preserva a hierarquia de headings
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                prefix = "#" * int(level) if level.isdigit() else "##"
                text_parts.append(f"\n{prefix} {text}\n")
            else:
                text_parts.append(text)

        # --- Tabelas ---
        for table in doc.tables:
            metadata["tables"] += 1
            table_text = self._format_table(table)
            text_parts.append(table_text)

        return ExtractedContent(
            text="\n".join(text_parts),
            metadata=metadata,
            file_type="docx",
        )

    @staticmethod
    def _format_table(table) -> str:
        """Converte tabela DOCX em texto legível."""
        lines = ["\n[Tabela]"]

        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
            if i == 0:
                lines.append("-" * 40)

        return "\n".join(lines)
